# from PIL import Image, ImageDraw
#
# BLACK = (0, 0, 0)
# YELLOW = (255, 255, 0)
# RED = (255, 0, 0)
# GREEN = (0, 128, 0)
# BLUE = (0, 255, 255)
# W, H = (200, 200)
# RW = 50
# RH = 50
#
# canvas = Image.new("RGB", (W, H), BLACK)  # создали холcт
# draw = ImageDraw.Draw(canvas)  # коробка с инструментами

# draw.rectangle((0, 0, 50, 50), fill=RED) # квадрат 3  по диагонали
# draw.rectangle((75, 75, 125, 125), fill=GREEN)
# draw.rectangle((W-RW, H-RH, W, H), fill=BLUE)
# canvas.save('image1.png', 'PNG')
# canvas.show()
# строим дугу

# from PIL import Image, ImageDraw
#
# BLACK = (0, 0, 0)
# YELLOW = (255, 255, 0)
# RED = (255, 0, 0)
# GREEN = (0, 255, 0)
# BLUE = (0, 255, 255) # светлый голубой
# W, H = (200, 200)
# RW = 50
# RH = 50
#
# canvas = Image.new("RGB", (W, H), RED)  # создали холcт
# draw = ImageDraw.Draw(canvas)  # коробка с инструментами
#
# draw.arc(xy=(25, 50, 175, 200), start=30, end=270, fill=BLUE, width=10)
#
# canvas.save('image2.png', 'PNG')
# canvas.show()

# рисуем многоугольник (треугольник)  .polygon()

# from PIL import Image, ImageDraw
#
# BLACK = (0, 0, 0)
# WHITE = (255, 255, 255)
# YELLOW = (255, 255, 0)
# RED = (255, 0, 0)
# GREEN = (0, 255, 0)
# BLUE = (0, 255, 255) # светлый голубой
# W, H = (200, 200)
# RW = 50
# RH = 50
#
# canvas = Image.new("RGB", (W, H), RED)  # создали холcт
# draw = ImageDraw.Draw(canvas)  # коробка с инструментамиRH = 50
#
# draw.polygon((W//2, 0, 0, H, W, H), fill=WHITE, outline=YELLOW)
# canvas.save('image2.png', 'PNG')
# canvas.show()

# from PIL import Image, ImageDraw
#
# BLACK = (0, 0, 0)
# WHITE = (255, 255, 255)
# YELLOW = (255, 255, 0)
# RED = (255, 0, 0)
# GREEN = (0, 255, 0)
# BLUE = (0, 255, 255) # светлый голубой
# W, H = (200, 200)
# RW = 50
# RH = 50
# T
# canvas = Image.new("RGB", (W, H), BLACK)  # создали холcт
# draw = ImageDraw.Draw(canvas)  # коробка с инструментамиRH = 50
#
# draw.ellipse((150, 0, 200, 50), fill=RED) # квадрат 3  по диагонали
# # draw.ellipse((150 = 200-50 x1, 0 - y1, 200 - x2, 50 - y2)
# draw.ellipse((75, 75, 125, 125), fill=GREEN)
# draw.ellipse((0, 150, 50, 200), fill=BLUE)
#
# canvas.show()

# РУСЕМ ТЕКСТ

# from PIL import Image, ImageDraw
# from PIL import ImageFont
#
# BLACK = (0, 0, 0)
# WHITE = (255, 255, 255)
# YELLOW = (255, 255, 0)
# RED = (255, 0, 0)
# GREEN = (0, 255, 0)
# BLUE = (0, 255, 255) # светлый голубой
# W, H = (200, 200)
# RW = 50
# RH = 50
# TEXT = 'PYTHON'
#
# Font = ImageFont.truetype('arial.ttf', size=20)
#
# canvas = Image.new("RGB", (W, H), BLACK)  # создали холcт
# draw = ImageDraw.Draw(canvas)
# draw.text((10, 10), TEXT, Font=Font)

# canvas.show()

# РИсуем Шахматная доска
# from PIL import Image, ImageDraw
# from PIL import ImageFont
#
#
# BLACK = (0, 0, 0)
# WHITE = (255, 255, 255)
# YELLOW = (255, 255, 0)
# RED = (255, 0, 0)
# GREEN = (0, 255, 0)
# BLUE = (0, 255, 255) # светлый голубой
# RW = 50
# RH = 50
# TEXT = 'PYTHON'
# CELLS = 8
# W, H = CELLS * RW, CELLS * RH
#
# canvas = Image.new("RGB", (W, H), BLACK)  # создали холcт
# draw = ImageDraw.Draw(canvas)
#
# for x in range(CELLS):
#         for y in range((x+1) % 2, CELLS, 2):
#             draw.rectangle((x * RW, y * RW,
#                             (x + 1) * RW - 1,
#                             (y + 1) * RW - 1), fill = WHITE )
# canvas.show()


# from PIL import Image, ImageDraw
# from PIL import ImageFont
# # РИсуем
#
# BLACK = (0, 0, 0)
# WHITE = (255, 255, 255)
# YELLOW = (255, 255, 0)
# RED = (255, 0, 0)
# GREEN = (0, 255, 0)
# BLUE = (0, 255, 255) # светлый голубой
# RW = 50
# RH = 50
# TEXT = 'PYTHON'
# CELLS = 8
# original = Image.open('pitone.jpg')
# resize = original.resize((400, 300))
#
#
# print('Формат:', original.format)
# print('Цветовая схема:', original.mode)
# print('Размеры:', original.size)
# W, H = original.size
# required_height = 100 # обязательно хранить с высотой 100
# ratio = W / H # считаем пропорцию
# print('Размеры:', resize.size)
#
# resize = original.resize((int(required_height * ratio), required_height))
# resize.save('resize.JPEG')
#
# temp = Image.open('resize.JPEG')
# print(temp.format)
#
# cropped_image = original.crop(
#     (550, 0, 800, 200)
# )
# cropped_image.save('cropped.JPEG')
# pixels = original.load()
# print(pixels)

# for x in range(W): # меняем цвета
#     for y in range(H):
#         r, g, b = pixels[x,y]
#         pixels[x,y] =  g, b,r

# for x in range(W): # меняем цвета на оттенки серого
#     for y in range(H):
#         r, g, b = pixels[x,y]
#         average = (r+ g +b) //3
#         pixels[x,y] =  average, average, average
#
#
# original.save('grayscale.JPEG')

# for x in range(W // 2): # меняем ориентацию картинки (зеркальное отражение по вертикали)
#     for y in range(H):
#         r, g, b = pixels[x,y]
#         average = (r+ g +b) // 3
#         pixels[x,y], pixels[W-x-1, y] = pixels[W-x-1, y], pixels[x,y],
#
# original.save('h_flipped.JPEG')
# original.save('h_flipped.JPEG')

# pixels = original.load()

# mode = original.mode
# original = original.convert('RGB')


# cropped_image.show()
# resize.show() # Размеры: (127, 100)

import numpy as np
from PIL import Image, ImageFilter, ImageEnhance
# BLACK = (0, 0, 0)
# WHITE = (255, 255, 255)
# YELLOW = (255, 255, 0)
# RED = (255, 0, 0)
# GREEN = (0, 255, 0)
# BLUE = (0, 255, 255) # светлый голубой
# RW = 50
# RH = 50
# TEXT = 'PYTHON'
# CELLS = 8

# original = Image.open('pitone.jpg')
# original = original.convert('RGB')
# W, H = original.size
# pixels = original.load()
# print(pixels)
# blur = original.filter(ImageFilter.BLUR) # размытие легкое
# blur.save('pitone_blur.JPEG')
#
#
# for x in range(W): # меняем цвета на оттенки серого
#     for y in range(H):
#         r, g, b = pixels[x,y]
#         average = (r+ g +b) //3
#         pixels[x,y] =  average, average, average




# blurbox = original.filter(ImageFilter.BoxBlur(20))  # размытие по параметру BoxBlur(20)
# blurbox.save('pitone_box_blur.JPEG')

# var = blur.filter(ImageFilter.FIND_EDGES)
# # var.show()
#
# var2 = np.asarray(var) + 255
# Image.fromarray(var2).show()


# blurgauss = original.filter(ImageFilter.GaussianBlur(20))
# blurgauss.save('pitone_blurgauss.JPEG')

# contour = original.filter(ImageFilter.CONTOUR())  # контур
# contour.save('pitone_contour.JPEG')

# делаем резкoсть


# C:\Users\Student\AppData\Local\Programs\Python\Python39\Scripts\pip.exe
# sharped = ImageEnhance.Sharpness(original)
# sharped_image = sharped.enhance(10.0)
# sharped_image.save('sharped.JPEG')



# переменая окружения

# from pymorphy2 import MorphAnalyzer
#
# form = MorphAnalyzer().parse('бутылка') [0]
#
# for bottle in reversed(range(9)):     # обратный отсчет от 9 до 0
#     print(f'В холодильнике {bottle +1} {form.make_agree_with_number(bottle + 1).word} пива')
#     print('Ввозьмем одну и выпьем.')
#
#     if bottle % 10 == 1 and bottle != 11:
#         remain = 'Осталась'
#     else:
#         remain = 'Осталось'
# print(f'{remain} {bottle} {form.make_agree_with_number(bottle + 1).word} пива')

from docx import Document

document =Document()

document.add_heading('Не забыть взять с собой', 0)
document.add_paragraph('Паспорт, удостоверение')

p = document.add_paragraph('Также надо ')
p.add_run('не забыть, ').bold = True

document.add_paragraph('это будет булетом', style='List Bullet')
document.add_paragraph('щетка', style='List Number')
document.add_paragraph('паста', style='List Number')

document.add_paragraph('"Варкалосьб хлипкие шарки пыряли по нове и хрукатали зелюки как мюмзики в нове"', style= 'Intense Quote')


table = document.add_table(rows=1,cols=2)
header_cells = table.rows[0].cells
header_cells[0].text = '№'
header_cells[1].text = 'Количество'


document.save('text.docx')

